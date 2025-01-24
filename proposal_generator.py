from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import RGBColor
import csv
from datetime import datetime

class ProposalGenerator:
    def __init__(self, student_name, roll_number):
        self.student_name = student_name
        self.roll_number = roll_number
        self.doc = Document()
        self.setup_document_formatting()
        self.generate_proposal()

    def setup_document_formatting(self):
        # Set up the default style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        # Set 1.5 line spacing
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_heading_with_format(self, text, level=1):
        heading = self.doc.add_heading(text, level=level)
        heading.style.font.name = 'Times New Roman'
        heading.style.font.size = Pt(12)
        return heading

    def add_section_with_word_limit(self, content, word_limit):
        words = content.split()
        if len(words) > word_limit:
            raise ValueError(f"Content exceeds word limit of {word_limit} words")
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run(content)
        return para

    def generate_proposal(self):
        self.add_title_page()
        self.add_declaration()
        self.add_executive_summary()
        self.add_organization_background()
        self.add_problem_statement()
        self.add_problem_background()
        self.add_problem_solving_approach()
        self.add_timeline()
        self.add_expected_outcome()

    def add_title_page(self):
        title = "Optimizing Inventory Management and Sales Analysis for Kirana Store Operations"
        self.doc.add_heading(title, 0)
        self.doc.add_paragraph()
        self.doc.add_paragraph(f"Submitted by:\n{self.student_name}\nRoll Number: {self.roll_number}")
        self.doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        self.doc.add_page_break()

    def add_declaration(self):
        self.add_heading_with_format("Declaration")
        declaration = (
            f"I, {self.student_name} (Roll No: {self.roll_number}), hereby declare that this "
            "project proposal is my original work and contains no plagiarized content. "
            "The plagiarism level is maintained below 20% as per the institutional requirements. "
            "All sources used have been properly cited and acknowledged."
        )
        self.add_section_with_word_limit(declaration, 100)
        self.doc.add_page_break()

    def add_executive_summary(self):
        self.add_heading_with_format("Executive Summary")
        summary = (
            "This project proposal aims to optimize the operational efficiency of a Kirana store "
            "through comprehensive data analysis and management solutions. The study focuses on "
            "analyzing transaction patterns, inventory management, and customer behavior to identify "
            "key areas for improvement. By leveraging existing transaction data and implementing "
            "advanced analytical tools, this project will provide actionable insights for better "
            "business decision-making. The proposed solutions will address inventory optimization, "
            "sales forecasting, and customer relationship management, ultimately leading to "
            "increased profitability and operational efficiency."
        )
        self.add_section_with_word_limit(summary, 250)

    def add_organization_background(self):
        self.add_heading_with_format("Organization Background")
        background = (
            "The Kirana store under study is a traditional Indian retail establishment that has "
            "been serving local communities for several years. Operating in a competitive market, "
            "the store offers a wide range of daily necessities, groceries, and household items. "
            "With increasing competition from modern retail formats and online platforms, the store "
            "seeks to modernize its operations while maintaining its traditional customer-centric "
            "approach. The business currently manages inventory manually and tracks sales through "
            "basic digital records, presenting opportunities for optimization through data-driven "
            "decision-making."
        )
        self.add_section_with_word_limit(background, 200)

    def add_problem_statement(self):
        self.add_heading_with_format("Problem Statement")
        self.doc.add_paragraph("The key objectives of this project are:", style='List Bullet')
        objectives = [
            "To analyze and optimize the current inventory management system, reducing stockouts and overstock situations",
            "To develop a data-driven approach for sales forecasting and demand prediction",
            "To identify patterns in customer purchasing behavior for improved product placement and marketing strategies"
        ]
        for objective in objectives:
            self.doc.add_paragraph(objective, style='List Bullet')

    def add_problem_background(self):
        self.add_heading_with_format("Background of the Problem")
        background = ("Traditional Kirana stores face significant challenges in the modern retail landscape. "
                     "Manual inventory management often leads to inefficiencies, while lack of data analysis "
                     "prevents optimal decision-making. The store's current system lacks proper tracking of "
                     "stock levels, leading to frequent stockouts or excess inventory. Additionally, "
                     "understanding customer preferences and purchasing patterns remains intuitive rather "
                     "than data-driven, limiting the store's ability to optimize its product mix and "
                     "marketing strategies.")
        self.add_section_with_word_limit(background, 250)

    def add_problem_solving_approach(self):
        self.add_heading_with_format("Problem Solving Approach")
        
        # Methods section
        self.add_heading_with_format("Methods Used", level=2)
        methods = (
            "The project will employ quantitative analysis methods including:\n"
            "• Time series analysis for sales forecasting\n"
            "• Statistical analysis of inventory turnover rates\n"
            "• Pattern recognition in customer purchase behavior\n"
            "• Predictive modeling for demand forecasting"
        )
        self.doc.add_paragraph(methods)

        # Data collection section
        self.add_heading_with_format("Data Collection", level=2)
        data_collection = (
            "Data will be collected from multiple sources:\n"
            "• Historical transaction records from the store's database\n"
            "• Inventory records including stock levels and ordering patterns\n"
            "• Customer purchase history and frequency\n"
            "• Product categorization and pricing information"
        )
        self.doc.add_paragraph(data_collection)

        # Analysis tools section
        self.add_heading_with_format("Analysis Tools", level=2)
        tools = (
            "The following tools will be utilized:\n"
            "• Python for data processing and analysis\n"
            "• Pandas for data manipulation and cleaning\n"
            "• Matplotlib and Seaborn for data visualization\n"
            "• Excel for reporting and dashboard creation"
        )
        self.doc.add_paragraph(tools)

    def add_timeline(self):
        self.add_heading_with_format("Expected Timeline")
        # Create a simple table for the timeline
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Phase'
        header_cells[1].text = 'Activities'
        header_cells[2].text = 'Duration'

        # Add timeline entries
        timeline_data = [
            ('Phase 1', 'Data Collection and Cleaning', '2 weeks'),
            ('Phase 2', 'Data Analysis and Pattern Recognition', '3 weeks'),
            ('Phase 3', 'Model Development and Testing', '4 weeks'),
            ('Phase 4', 'Implementation and Documentation', '3 weeks')
        ]

        for phase, activities, duration in timeline_data:
            row_cells = table.add_row().cells
            row_cells[0].text = phase
            row_cells[1].text = activities
            row_cells[2].text = duration

    def add_expected_outcome(self):
        self.add_heading_with_format("Expected Outcome")
        outcome = (
            "The project is expected to deliver:\n"
            "• A comprehensive analysis of current inventory management practices\n"
            "• Data-driven recommendations for stock optimization\n"
            "• Predictive models for sales forecasting\n"
            "• Actionable insights for improved customer service and satisfaction\n"
            "• Documentation and training materials for sustainable implementation"
        )
        self.doc.add_paragraph(outcome)

    def save(self, filename):
        self.doc.save(filename)

def main():
    try:
        generator = ProposalGenerator("Chirantan Dey", "23f2001382")
        generator.save("Business_Proposal.docx")
        print("Proposal generated successfully!")
    except Exception as e:
        print(f"Error generating proposal: {str(e)}")

if __name__ == "__main__":
    main()